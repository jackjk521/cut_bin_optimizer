from fastapi import FastAPI, File, UploadFile
from fastapi.responses import FileResponse
import pandas as pd
from docx import Document
import os
import uuid
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi import Request

from docx.shared import Pt  # Import for setting font size


app = FastAPI()

# Enable CORS for all origins to allow automatic downloads
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve static files and templates
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

UPLOAD_DIR = "uploads"
OUTPUT_DIR = "output"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# def fill_contract(template_path, data_row):
    # doc = Document(template_path)
    # for para in doc.paragraphs:
    #     # Combine all runs into a single text
    #     full_text = "".join([run.text for run in para.runs])
    #     # Replace placeholders in the combined text
    #     for key, value in data_row.items():
    #         placeholder = f"{{{{{key.strip()}}}}}"
    #         if placeholder in full_text:
    #             full_text = full_text.replace(placeholder, str(value))
    #             # Formatting of updated fields
    #             full_text.capitalize
    #     # Clear existing runs and replace with updated text as a single run
    #     if full_text != "".join([run.text for run in para.runs]):
    #         for run in para.runs:
    #             run.text = ""  # Clear current text
    #         new_run = para.add_run(full_text)
    #         # new_run.bold = True  # Make filled text bold
    # filled_docx = os.path.join(OUTPUT_DIR, f"filled_{uuid.uuid4()}.docx")
    # doc.save(filled_docx)
    # return filled_docx


def fill_contract(template_path, data_row):
    doc = Document(template_path)
    for para in doc.paragraphs:
        # Combine all runs into a single text string
        full_text = "".join([run.text for run in para.runs])
        # Replace placeholders in the combined text
        for key, value in data_row.items():
            placeholder = f"{{{{{key.strip()}}}}}"
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(value).upper())

        # If replacements were made, update the paragraph runs
        if full_text != "".join([run.text for run in para.runs]):
            for run in para.runs:
                run.text = ""  # Clear current text
            # Add updated text as a single run with bold and font size 12
            new_run = para.add_run(full_text)
            new_run.bold = True
            new_run.font.size = Pt(12)

    # Handle placeholders inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text = "".join([run.text for run in para.runs])
                    for key, value in data_row.items():
                        placeholder = f"{{{{{key.strip()}}}}}"
                        if placeholder in full_text:
                            full_text = full_text.replace(placeholder, str(value).upper())
                    if full_text != "".join([run.text for run in para.runs]):
                        for run in para.runs:
                            run.text = ""
                        # Add updated text as a single run with bold and font size 12
                        new_run = para.add_run(full_text)
                        new_run.bold = True
                        new_run.font.size = Pt(12)

    filled_docx = os.path.join(OUTPUT_DIR, f"filled_{uuid.uuid4()}.docx")
    doc.save(filled_docx)
    return filled_docx


@app.get("/")
async def read_index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/upload")
async def upload_files(csv_file: UploadFile = File(...), template: UploadFile = File(...)):
    csv_path = os.path.join(UPLOAD_DIR, csv_file.filename)
    template_path = os.path.join(UPLOAD_DIR, template.filename)
    with open(csv_path, "wb") as f:
        f.write(await csv_file.read())
    with open(template_path, "wb") as f:
        f.write(await template.read())

    df = pd.read_csv(csv_path)

    # Preprocess columns: fill empty or null values using direct assignment
    df['BUILDING'] = df['BUILDING'].fillna(-1).astype(str)
    df['UNIT/SLOT'] = df['UNIT/SLOT'].fillna(-1).astype(str)
    df['AREA_(SQM)'] = df['AREA_(SQM)'].fillna(-1)
    df['BUSINESS_NAME'] = df['BUSINESS_NAME'].fillna('MISSING BUSINESS NAME')
    df['ADDRESS'] = df['ADDRESS'].fillna('BAGUIO CITY')

    # Create a new column for "BUILDING UNIT" as string
    df['BUILDING_UNIT'] = df.apply(lambda row: f"{str(row['BUILDING'])} {str(row['UNIT/SLOT'])}", axis=1)

    # Capitalize all text columns
    df = df.applymap(lambda x: x.upper() if isinstance(x, str) else x)
    # Save preprocessed CSV
    preprocessed_csv = os.path.join(UPLOAD_DIR, "preprocessed_" + csv_file.filename)
    df.to_csv(preprocessed_csv, index=False)

    docx_files = []
    for _, row in df.iterrows():
        docx_file = fill_contract(template_path, row)
        docx_files.append(docx_file)

    return FileResponse(preprocessed_csv, media_type='text/csv', filename=os.path.basename(preprocessed_csv))

@app.get("/download/{file_name}")
async def download_file(file_name: str):
    file_path = os.path.join(OUTPUT_DIR, file_name)
    return FileResponse(file_path)
